import json
import os
import re

import pdfplumber
from docx import Document


# ==================================================
# LOAD ALL SKILLS FROM jobs.json
# ==================================================

def load_skills():

    base_dir = os.path.dirname(
        os.path.abspath(__file__)
    )

    jobs_file = os.path.join(
        base_dir,
        "data",
        "jobs.json"
    )

    with open(
        jobs_file,
        "r",
        encoding="utf-8"
    ) as file:

        jobs_data = json.load(file)

    skills = []

    for job_data in jobs_data.values():

        for skill in job_data.get("skills", []):

            if skill not in skills:
                skills.append(skill)

    return skills


# ==================================================
# SKILL ALIASES
# ==================================================
# Resume users may write skills in different ways.
# Every alias maps back to the official skill name
# used inside jobs.json.

SKILL_ALIASES = {

    # ---------------- PROGRAMMING ----------------

    "python3": "Python",
    "python 3": "Python",

    "js": "JavaScript",
    "javascript": "JavaScript",

    "cpp": "C++",
    "c plus plus": "C++",

    # ---------------- WEB ----------------

    "reactjs": "React",
    "react.js": "React",
    "react js": "React",

    "nodejs": "Node.js",
    "node js": "Node.js",
    "node.js": "Node.js",

    "expressjs": "Express.js",
    "express js": "Express.js",
    "express.js": "Express.js",

    "rest api": "REST API",
    "rest apis": "REST API",
    "restful api": "REST API",
    "restful apis": "REST API",

    # ---------------- DATA ----------------

    "powerbi": "Power BI",
    "power bi": "Power BI",

    "ms excel": "Excel",
    "microsoft excel": "Excel",

    "numpy": "NumPy",

    "pandas": "Pandas",

    "sklearn": "Scikit-learn",
    "scikit learn": "Scikit-learn",
    "scikit-learn": "Scikit-learn",

    # ---------------- AI / ML ----------------

    "ml": "Machine Learning",
    "machine learning": "Machine Learning",

    "dl": "Deep Learning",
    "deep learning": "Deep Learning",

    "tensorflow": "TensorFlow",

    "pytorch": "PyTorch",

    "nlp": "Natural Language Processing",
    "natural language processing":
        "Natural Language Processing",

    "cv": "Computer Vision",
    "computer vision": "Computer Vision",

    # ---------------- BIG DATA ----------------

    "spark": "Apache Spark",
    "apache spark": "Apache Spark",

    "apache airflow": "Airflow",

    # ---------------- DATABASE ----------------

    "mongodb": "MongoDB",
    "mongo db": "MongoDB",

    "sqlite": "SQLite",

    # ---------------- CLOUD ----------------

    "amazon web services": "AWS",
    "aws": "AWS",

    "microsoft azure": "Azure",

    # ---------------- DEVOPS ----------------

    "k8s": "Kubernetes",

    "ci cd": "CI/CD",
    "ci/cd": "CI/CD",
    "cicd": "CI/CD",

    # ---------------- VERSION CONTROL ----------------

    "github": "Git",
    "gitlab": "Git",

    # ---------------- TESTING ----------------

    "selenium webdriver": "Selenium",

    "postman api": "Postman",

    "jira": "JIRA",

    # ---------------- ANDROID ----------------

    "android sdk": "Android Studio",

    # ---------------- IOS ----------------

    "swift ui": "SwiftUI",

    "ios sdk": "iOS SDK"
}


# ==================================================
# EXTRACT TEXT FROM RESUME
# ==================================================

def extract_text(file_path):

    text = ""

    file_path_lower = file_path.lower()

    # ------------------------------------------------
    # PDF
    # ------------------------------------------------

    if file_path_lower.endswith(".pdf"):

        with pdfplumber.open(file_path) as pdf:

            for page in pdf.pages:

                page_text = page.extract_text()

                if page_text:

                    text += page_text + "\n"

    # ------------------------------------------------
    # DOCX
    # ------------------------------------------------

    elif file_path_lower.endswith(".docx"):

        document = Document(file_path)

        for paragraph in document.paragraphs:

            text += paragraph.text + "\n"

        # Also read text inside tables
        for table in document.tables:

            for row in table.rows:

                for cell in row.cells:

                    text += cell.text + "\n"

    # ------------------------------------------------
    # UNSUPPORTED FILE
    # ------------------------------------------------

    else:

        raise ValueError(
            "Unsupported resume format. "
            "Use PDF or DOCX."
        )

    return text


# ==================================================
# NORMALIZE TEXT
# ==================================================

def normalize_text(text):

    text = text.lower()

    # Convert line breaks/tabs into spaces
    text = text.replace("\n", " ")
    text = text.replace("\t", " ")

    # Remove repeated spaces
    text = re.sub(
        r"\s+",
        " ",
        text
    )

    return text.strip()


# ==================================================
# CHECK WHETHER A SKILL EXISTS
# ==================================================

def skill_exists(skill_text, resume_text):

    skill_text = skill_text.lower().strip()

    # Escape special regex characters such as:
    # C++, Node.js, CI/CD etc.
    escaped_skill = re.escape(skill_text)

    pattern = (
        r"(?<!\w)"
        + escaped_skill
        + r"(?!\w)"
    )

    return bool(
        re.search(
            pattern,
            resume_text,
            flags=re.IGNORECASE
        )
    )


# ==================================================
# EXTRACT SKILLS FROM RESUME
# ==================================================

def extract_skills(file_path):

    # Extract resume text
    resume_text = extract_text(
        file_path
    )

    # Normalize text
    resume_text = normalize_text(
        resume_text
    )

    # Load official skills from jobs.json
    skills = load_skills()

    found_skills = []

    # ------------------------------------------------
    # 1. EXACT OFFICIAL SKILL MATCHING
    # ------------------------------------------------

    for skill in skills:

        if skill_exists(
            skill,
            resume_text
        ):

            if skill not in found_skills:

                found_skills.append(
                    skill
                )

    # ------------------------------------------------
    # 2. ALIAS MATCHING
    # ------------------------------------------------

    for alias, official_skill in (
        SKILL_ALIASES.items()
    ):

        if skill_exists(
            alias,
            resume_text
        ):

            # Only return skills that are actually
            # supported by jobs.json
            if (
                official_skill in skills
                and official_skill
                not in found_skills
            ):

                found_skills.append(
                    official_skill
                )

    # ------------------------------------------------
    # RETURN DETECTED SKILLS
    # ------------------------------------------------

    return found_skills